#!/usr/bin/env python3
"""
Enhanced Resume Parser with AI Integration
Combines PyResParser with advanced AI analysis
"""

import os
import re
import json
from typing import Dict, List, Optional
from collections import defaultdict

# Resume parsing
try:
    from pyresparser import ResumeParser
    PYRESPARSER_AVAILABLE = True
except ImportError:
    print("⚠️ PyResParser not available, using AI-only extraction")
    PYRESPARSER_AVAILABLE = False

# AI Model
from advanced_ai_model import AdvancedResumeAI

class EnhancedResumeParser:
    """Enhanced Resume Parser with AI Integration"""
    
    def __init__(self):
        """Initialize the enhanced parser"""
        print("🚀 Initializing Enhanced Resume Parser...")
        
        # Initialize AI model
        self.ai_model = AdvancedResumeAI()
        
        # Initialize skill requirements for different roles
        self.role_skill_requirements = {
            'Software Engineer': {
                'required': ['python', 'java', 'javascript', 'git', 'sql'],
                'preferred': ['react', 'node.js', 'aws', 'docker', 'kubernetes']
            },
            'Data Scientist': {
                'required': ['python', 'sql', 'machine learning', 'pandas', 'numpy'],
                'preferred': ['tensorflow', 'pytorch', 'spark', 'tableau', 'r']
            },
            'Product Manager': {
                'required': ['project management', 'agile', 'scrum', 'analytics'],
                'preferred': ['jira', 'confluence', 'sql', 'tableau', 'figma']
            },
            'DevOps Engineer': {
                'required': ['docker', 'kubernetes', 'aws', 'jenkins', 'git'],
                'preferred': ['terraform', 'ansible', 'prometheus', 'grafana']
            }
        }
        
        print("✅ Enhanced Resume Parser initialized!")
    
    def parse_resume_with_pyresparser(self, file_path: str) -> Dict:
        """Extract structured data using PyResParser with AI fallback"""
        try:
            print(f"📄 Parsing resume: {file_path}")
            
            if not PYRESPARSER_AVAILABLE:
                return self._ai_based_extraction(file_path)
            
            # Use PyResParser to extract structured data
            data = ResumeParser(file_path).get_extracted_data()
            
            # Clean and enhance the extracted data
            enhanced_data = {
                'name': self._clean_name(data.get('name')),
                'email': self._clean_email(data.get('email')),
                'mobile_number': self._clean_mobile(data.get('mobile_number')),
                'skills': self._clean_skills(data.get('skills', [])),
                'education': self._clean_education(data.get('education', [])),
                'experience': self._clean_experience(data.get('experience', [])),
                'total_experience': data.get('total_experience', 0),
                'college_name': data.get('college_name', []),
                'degree': data.get('degree', []),
                'designation': data.get('designation', []),
                'company_names': data.get('company_names', [])
            }
            
            return enhanced_data
            
        except Exception as e:
            print(f"❌ Error parsing resume with PyResParser: {e}")
            print("🔄 Using AI-based extraction as fallback...")
            return self._ai_based_extraction(file_path)
    
    def _ai_based_extraction(self, file_path: str) -> Dict:
        """AI-based extraction as fallback when PyResParser fails"""
        try:
            # Extract text from file
            if file_path.endswith('.pdf'):
                resume_text = self.ai_model.extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    resume_text = f.read()
            
            if not resume_text:
                return self._get_empty_data()
            
            # Use AI model to extract features
            features = self.ai_model.extract_advanced_features(resume_text)
            
            # Extract name using regex patterns
            name_patterns = [
                r'^([A-Z][a-z]+ [A-Z][a-z]+)',  # First Last
                r'Name[:\s]+([A-Z][a-z]+ [A-Z][a-z]+)',  # Name: First Last
                r'([A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+)',  # First M. Last
            ]
            
            name = "Not specified"
            for pattern in name_patterns:
                match = re.search(pattern, resume_text)
                if match:
                    name = match.group(1).strip()
                    break
            
            # Extract email
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            email_match = re.search(email_pattern, resume_text)
            email = email_match.group(0) if email_match else "Not specified"
            
            # Extract phone
            phone_patterns = [
                r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
                r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
                r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'
            ]
            
            phone = "Not specified"
            for pattern in phone_patterns:
                match = re.search(pattern, resume_text)
                if match:
                    phone = match.group(0)
                    break
            
            # Extract skills using AI model's skill taxonomy
            skills = []
            for category, skill_set in self.ai_model.skill_taxonomy.items():
                found_skills = features.get(f'{category}_skills', [])
                skills.extend(found_skills)
            
            # Extract experience years
            experience_years = features.get('experience_years', 0)
            
            # Extract education and companies from text
            education = self._extract_education_from_text(resume_text)
            companies = self._extract_companies_from_text(resume_text)
            degrees = self._extract_degrees_from_text(resume_text)
            
            return {
                'name': name,
                'email': email,
                'mobile_number': phone,
                'skills': skills[:20],  # Limit to top 20 skills
                'education': education,
                'experience': [],
                'total_experience': experience_years,
                'college_name': [],
                'degree': degrees,
                'designation': [],
                'company_names': companies
            }
            
        except Exception as e:
            print(f"❌ AI-based extraction failed: {e}")
            return self._get_empty_data()
    
    def _get_empty_data(self) -> Dict:
        """Return empty data structure"""
        return {
            'name': "Not specified",
            'email': "Not specified", 
            'mobile_number': "Not specified",
            'skills': [],
            'education': [],
            'experience': [],
            'total_experience': 0,
            'college_name': [],
            'degree': [],
            'designation': [],
            'company_names': []
        }
    
    def _clean_name(self, name) -> str:
        """Clean and format the extracted name"""
        if not name:
            return "Not specified"
        
        if isinstance(name, list):
            name = ' '.join(name)
        
        # Remove common prefixes and clean
        name = str(name).strip()
        name = re.sub(r'^(Mr\.?|Ms\.?|Mrs\.?|Dr\.?)\s*', '', name, flags=re.IGNORECASE)
        
        return name.title() if name else "Not specified"
    
    def _clean_email(self, email) -> str:
        """Clean and validate email"""
        if not email:
            return "Not specified"
        
        if isinstance(email, list):
            email = email[0] if email else ""
        
        email = str(email).strip().lower()
        
        # Basic email validation
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if re.match(email_pattern, email):
            return email
        
        return "Not specified"
    
    def _clean_mobile(self, mobile) -> str:
        """Clean and format mobile number"""
        if not mobile:
            return "Not specified"
        
        if isinstance(mobile, list):
            mobile = mobile[0] if mobile else ""
        
        # Clean mobile number
        mobile = str(mobile).strip()
        mobile = re.sub(r'[^\d\+\-\(\)\s]', '', mobile)
        
        return mobile if mobile else "Not specified"
    
    def _clean_skills(self, skills) -> List[str]:
        """Clean and enhance skills list"""
        if not skills:
            return []
        
        if isinstance(skills, str):
            skills = [skills]
        
        cleaned_skills = []
        for skill in skills:
            if skill and isinstance(skill, str):
                skill = skill.strip().lower()
                if len(skill) > 1:
                    cleaned_skills.append(skill)
        
        return list(set(cleaned_skills))  # Remove duplicates
    
    def _clean_education(self, education) -> List[str]:
        """Clean education information"""
        if not education:
            return []
        
        if isinstance(education, str):
            return [education]
        
        return [str(edu).strip() for edu in education if edu]
    
    def _clean_experience(self, experience) -> List[str]:
        """Clean experience information"""
        if not experience:
            return []
        
        if isinstance(experience, str):
            return [experience]
        
        return [str(exp).strip() for exp in experience if exp]
    
    def _extract_education_from_text(self, text: str) -> List[str]:
        """Extract education information from text"""
        education_keywords = ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd', 'degree']
        education = []
        
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in education_keywords):
                if len(line.strip()) > 10:  # Avoid short matches
                    education.append(line.strip())
        
        return education[:3]  # Limit to 3 entries
    
    def _extract_companies_from_text(self, text: str) -> List[str]:
        """Extract company names from text"""
        company_indicators = ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems', 'solutions']
        companies = []
        
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            if any(indicator in line_lower for indicator in company_indicators):
                if len(line.strip()) > 5:
                    companies.append(line.strip())
        
        return companies[:5]  # Limit to 5 companies
    
    def _extract_degrees_from_text(self, text: str) -> List[str]:
        """Extract degree information from text"""
        degree_patterns = [
            r'\b(Bachelor|Master|PhD|Ph\.D|MBA|MS|BS|BA|MA|B\.Tech|M\.Tech)\b',
            r'\b(B\.S\.|M\.S\.|B\.A\.|M\.A\.)\b'
        ]
        
        degrees = []
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            degrees.extend(matches)
        
        return list(set(degrees))  # Remove duplicates
    
    def comprehensive_resume_analysis(self, file_path: str, target_role: str = None) -> Dict:
        """Perform comprehensive resume analysis combining all features"""
        
        print(f"🔍 Starting comprehensive analysis for: {file_path}")
        
        # Extract text for AI analysis
        if file_path.endswith('.pdf'):
            resume_text = self.ai_model.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                resume_text = f.read()
        
        # Extract structured data using PyResParser
        extracted_data = self.parse_resume_with_pyresparser(file_path)
        
        # Get AI-powered analysis
        ai_analysis = self.ai_model.comprehensive_analysis(resume_text)
        
        # Predict suitable roles with enhanced matching
        role_predictions = self.predict_suitable_roles(resume_text, extracted_data)
        
        # Generate skill recommendations
        skill_recommendations = self.generate_skill_recommendations(extracted_data, target_role)
        
        # Generate enhancement suggestions
        enhancement_suggestions = self.generate_resume_enhancement_suggestions(extracted_data, resume_text)
        
        # Determine gender pronoun (simple heuristic based on name)
        name = extracted_data.get('name', '')
        gender_pronoun = self._guess_gender_pronoun(name)
        
        # Compile comprehensive result
        result = {
            'extracted_data': extracted_data,
            'ai_analysis': ai_analysis,
            'role_predictions': role_predictions,
            'skill_recommendations': skill_recommendations,
            'enhancement_suggestions': enhancement_suggestions,
            'gender_pronoun': gender_pronoun,
            'analysis_summary': {
                'name': extracted_data.get('name', 'Not specified'),
                'email': extracted_data.get('email', 'Not specified'),
                'phone': extracted_data.get('mobile_number', 'Not specified'),
                'total_skills': len(extracted_data.get('skills', [])),
                'experience_years': extracted_data.get('total_experience', 0),
                'top_role_prediction': role_predictions[0] if role_predictions else None,
                'resume_quality_score': ai_analysis.get('quality_assessment', {}).get('quality_score', 0)
            }
        }
        
        print("✅ Comprehensive analysis completed!")
        return result
    
    def predict_suitable_roles(self, resume_text: str, extracted_data: Dict) -> List[Dict]:
        """Predict suitable roles using AI with skill matching enhancement"""
        
        # Get AI predictions
        ai_predictions = self.ai_model.predict_job_roles(resume_text, top_k=8)
        
        # Enhance predictions with skill matching
        enhanced_predictions = []
        user_skills = set(skill.lower() for skill in extracted_data.get('skills', []))
        
        for prediction in ai_predictions:
            role_name = prediction['role']
            
            # Find matching role in our skill requirements
            matching_role = None
            for role_key in self.role_skill_requirements:
                if role_key.lower() in role_name.lower() or role_name.lower() in role_key.lower():
                    matching_role = role_key
                    break
            
            if matching_role:
                requirements = self.role_skill_requirements[matching_role]
                required_skills = set(skill.lower() for skill in requirements['required'])
                preferred_skills = set(skill.lower() for skill in requirements['preferred'])
                
                # Calculate skill match scores
                required_match = len(user_skills.intersection(required_skills))
                preferred_match = len(user_skills.intersection(preferred_skills))
                
                skill_score = (required_match / len(required_skills)) * 70 + (preferred_match / len(preferred_skills)) * 30
                
                # Adjust prediction probability based on skill match
                adjusted_probability = (prediction['probability'] + skill_score) / 2
                
                enhanced_predictions.append({
                    **prediction,
                    'adjusted_probability': round(adjusted_probability, 2),
                    'skill_match_score': round(skill_score, 2),
                    'required_skills_matched': required_match,
                    'total_required_skills': len(required_skills),
                    'preferred_skills_matched': preferred_match,
                    'total_preferred_skills': len(preferred_skills),
                    'missing_required_skills': list(required_skills - user_skills),
                    'missing_preferred_skills': list(preferred_skills - user_skills)
                })
            else:
                enhanced_predictions.append({
                    **prediction,
                    'adjusted_probability': prediction['probability'],
                    'skill_match_score': 0,
                    'missing_required_skills': [],
                    'missing_preferred_skills': []
                })
        
        # Sort by adjusted probability
        enhanced_predictions.sort(key=lambda x: x['adjusted_probability'], reverse=True)
        
        return enhanced_predictions
    
    def generate_skill_recommendations(self, extracted_data: Dict, target_role: str = None) -> Dict:
        """Generate personalized skill recommendations"""
        
        user_skills = set(skill.lower() for skill in extracted_data.get('skills', []))
        
        recommendations = {
            'immediate_skills': [],
            'trending_skills': ['python', 'aws', 'cybersecurity', 'data science', 'react', 'kubernetes'],
            'role_specific': {}
        }
        
        # Basic skill recommendations
        basic_skills = ['python', 'git', 'sql', 'java', 'javascript']
        recommendations['immediate_skills'] = [skill for skill in basic_skills if skill not in user_skills]
        
        # Role-specific recommendations
        if target_role:
            for role_key, requirements in self.role_skill_requirements.items():
                if target_role.lower() in role_key.lower():
                    missing_required = [skill for skill in requirements['required'] if skill not in user_skills]
                    missing_preferred = [skill for skill in requirements['preferred'] if skill not in user_skills]
                    
                    recommendations['role_specific'][role_key] = {
                        'missing_required': missing_required,
                        'missing_preferred': missing_preferred
                    }
                    break
        
        return recommendations
    
    def generate_resume_enhancement_suggestions(self, extracted_data: Dict, resume_text: str) -> List[str]:
        """Generate personalized resume enhancement suggestions"""
        
        suggestions = []
        
        # Check basic information completeness
        if extracted_data.get('name') == "Not specified":
            suggestions.append("🔧 Add your full name prominently at the top")
        
        if extracted_data.get('email') == "Not specified":
            suggestions.append("📧 Include a professional email address")
        
        if extracted_data.get('mobile_number') == "Not specified":
            suggestions.append("📱 Add your contact phone number")
        
        # Check skills
        skills_count = len(extracted_data.get('skills', []))
        if skills_count < 5:
            suggestions.append("🔧 Add more technical skills to strengthen your profile")
        
        # Check experience
        if extracted_data.get('total_experience', 0) == 0:
            suggestions.append("💼 Add detailed work experience with specific achievements")
        
        # Check education
        if not extracted_data.get('degree'):
            suggestions.append("🎓 Include your educational background")
        
        # Text analysis suggestions
        if len(resume_text) < 500:
            suggestions.append("📈 Expand your resume with more detailed descriptions")
        
        # Quality suggestions
        if 'achievement' not in resume_text.lower():
            suggestions.append("🏆 Add quantifiable achievements and accomplishments")
        
        if skills_count < 10:
            suggestions.append("📈 Overall resume quality needs improvement - consider restructuring")
        
        return suggestions[:8]  # Limit to top 8 suggestions
    
    def _guess_gender_pronoun(self, name: str) -> str:
        """Simple heuristic to guess gender pronoun based on name"""
        if not name or name == "Not specified":
            return "they/them"
        
        # Simple name-based heuristic (this is very basic and not always accurate)
        male_names = {'john', 'michael', 'david', 'james', 'robert', 'william', 'richard', 'thomas'}
        female_names = {'mary', 'patricia', 'jennifer', 'linda', 'elizabeth', 'barbara', 'susan', 'jessica'}
        
        first_name = name.split()[0].lower() if name else ""
        
        if first_name in male_names:
            return "he/him"
        elif first_name in female_names:
            return "she/her"
        else:
            return "they/them"  # Default to neutral